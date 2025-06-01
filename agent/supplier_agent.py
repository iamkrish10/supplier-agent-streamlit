import os
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import json
from langgraph.graph import StateGraph
from langchain_openai import ChatOpenAI
from typing import TypedDict
import io
import tempfile
from pathlib import Path
from .utils import (
    FileProcessor, DataValidator, AuditProcessor, 
    ChartGenerator, ScoreCalculator, ReportFormatter
)

class SupplierState(TypedDict):
    structured_data: list
    audit_findings: dict
    weights: dict
    score_table: dict
    summary: str
    graph_path: str
    uploaded_files: dict  # New: Store uploaded files

class SupplierAgent:
    def __init__(self, openai_api_key: str):
        self.openai_api_key = openai_api_key
        self.graph = self._build_graph()
        
    def _build_graph(self):
        """Build the LangGraph workflow"""
        builder = StateGraph(SupplierState)
        builder.add_node("load", self.load_inputs)
        builder.add_node("audit", self.extract_audit_findings)
        builder.add_node("estimate_weights", self.estimate_weights_with_llm)
        builder.add_node("score", self.calculate_score)
        builder.add_node("summarize", self.generate_summary)
        builder.add_node("chart", self.plot_graph)
        builder.add_node("report", self.final_output)

        builder.set_entry_point("load")
        builder.add_edge("load", "audit")
        builder.add_edge("audit", "estimate_weights")
        builder.add_edge("estimate_weights", "score")
        builder.add_edge("score", "summarize")
        builder.add_edge("summarize", "chart")
        builder.add_edge("chart", "report")
        builder.set_finish_point("report")
        
        return builder.compile()
    
    def process_supplier_data(self, excel_file, audit_files: list) -> dict:
        """
        Main entry point for Streamlit
        """
        # Initialize state with uploaded files
        initial_state = {
            "structured_data": [],
            "audit_findings": {},
            "weights": {},
            "score_table": {},
            "summary": "",
            "graph_path": "",
            "uploaded_files": {
                "excel": excel_file,
                "audits": audit_files
            }
        }
        
        try:
            # Run the graph
            final_state = self.graph.invoke(initial_state)
            return {
                "success": True,
                "data": final_state
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }

    def load_inputs(self, state: SupplierState) -> SupplierState:
        """Load data from uploaded Excel file"""
        try:
            excel_file = state["uploaded_files"]["excel"]
            
            # Use utility to read file
            df = FileProcessor.read_excel_file(excel_file)
            
            # Validate and clean data
            validation_results = DataValidator.validate_supplier_data(df)
            if not validation_results["is_valid"]:
                raise Exception(f"Data validation failed: {validation_results['errors']}")
            
            # Clean the data
            df_clean = DataValidator.clean_supplier_data(df)
            
            state["structured_data"] = df_clean.to_dict(orient="records")
            return state
            
        except Exception as e:
            raise Exception(f"Error loading supplier data: {str(e)}")

    def extract_audit_findings(self, state: SupplierState) -> SupplierState:
        """Extract audit findings from uploaded Word documents"""
        findings = {}
        audit_files = state["uploaded_files"]["audits"]
        
        for i, audit_file in enumerate(audit_files):
            try:
                # Save uploaded file temporarily using utility
                tmp_file_path = FileProcessor.save_temp_docx(audit_file)
                
                # Reset file pointer
                audit_file.seek(0)
                
                # Extract text using utility
                content = AuditProcessor.extract_text_from_docx(tmp_file_path)
                
                # Extract supplier name using utility
                supplier_name = AuditProcessor.extract_supplier_name(content, audit_file.name)
                
                # Count findings using utility
                count = AuditProcessor.count_audit_findings(content)
                findings[supplier_name] = count
                
                # Clean up temp file
                FileProcessor.cleanup_temp_file(tmp_file_path)
                
            except Exception as e:
                print(f"Error processing audit file {audit_file.name}: {str(e)}")
                continue
        
        state["audit_findings"] = findings
        return state

    def estimate_weights_with_llm(self, state: SupplierState) -> SupplierState:
        """Use LLM to estimate weights for scoring"""
        try:
            llm = ChatOpenAI(model="gpt-3.5-turbo", api_key=self.openai_api_key)

            prompt = (
                "As a supplier quality expert, recommend standard weightage percentages "
                "for calculating performance based on On-Time Delivery Rate, Defect Rate, and Audit Findings. "
                "Respond with a JSON dictionary where keys are 'OnTimeDeliveryRate', 'DefectRate', and 'AuditFindings'. "
                "The values should be decimal numbers that sum to 1.0."
            )
            
            result = llm.invoke(prompt)
            weights = json.loads(result.content.strip())
            state["weights"] = weights
            
        except Exception as e:
            print(f"Error with LLM: {str(e)}")
            # Use default weights
            state["weights"] = {
                "OnTimeDeliveryRate": 0.5,
                "DefectRate": 0.3,
                "AuditFindings": 0.2
            }
        
        return state

    def calculate_score(self, state: SupplierState) -> SupplierState:
        """Calculate supplier performance scores"""
        scores = {}
        max_audit_score = 5
        weights = state["weights"]

        for supplier in state["structured_data"]:
            name = supplier.get("Supplier")
            if not name:
                continue

            # Prepare metrics for calculation
            metrics = {
                "OnTimeDeliveryRate": supplier.get("OnTimeDeliveryRate", 0) / 100,
                "DefectRate": (100 - supplier.get("DefectRate", 0)) / 100,  # Invert defect rate
                "AuditFindings": (max_audit_score - state["audit_findings"].get(name, 0)) / max_audit_score
            }

            # Use utility to calculate weighted score
            score = ScoreCalculator.calculate_weighted_score(metrics, weights) * 100
            score = ScoreCalculator.normalize_score(score, 0, 100)
            
            scores[name] = round(score, 2)

        state["score_table"] = scores
        return state

    def generate_summary(self, state: SupplierState) -> SupplierState:
        """Generate summary using LLM"""
        try:
            llm = ChatOpenAI(model="gpt-3.5-turbo", api_key=self.openai_api_key)

            # Prepare data for LLM
            supplier_data = []
            for supplier in state["structured_data"]:
                name = supplier.get("Supplier", "Unknown")
                score = state["score_table"].get(name, 0)
                findings = state["audit_findings"].get(name, 0)
                on_time = supplier.get("OnTimeDeliveryRate", 0)
                defect = supplier.get("DefectRate", 0)
                
                # Add performance tier
                tier = ScoreCalculator.calculate_performance_tier(score)
                
                supplier_data.append(
                    f"{name}: Score {score} ({tier}), On-Time Delivery: {on_time}%, "
                    f"Defect Rate: {defect}%, Audit Findings: {findings}"
                )

            prompt = (
                "Analyze these supplier performance results and provide a comprehensive summary:\n\n"
                + "\n".join(supplier_data) + 
                "\n\nProvide insights on:\n"
                "1. Overall performance trends\n"
                "2. Top and bottom performers with specific reasons\n"
                "3. Key areas for improvement\n"
                "4. Risk assessment\n"
                "5. Actionable recommendations for supplier management\n\n"
                "Keep the summary professional and actionable for management decision-making."
            )
            
            result = llm.invoke(prompt)
            state["summary"] = result.content
            
        except Exception as e:
            print(f"Error generating summary: {str(e)}")
            # Fallback summary using utilities
            scores_df = ReportFormatter.format_score_table(state["score_table"])
            top_performer = scores_df.iloc[0] if not scores_df.empty else None
            bottom_performer = scores_df.iloc[-1] if not scores_df.empty else None
            
            fallback_summary = f"""
**Supplier Performance Analysis Summary**

**Total Suppliers Evaluated:** {len(state['score_table'])}

**Top Performer:** {top_performer['Supplier'] if top_performer is not None else 'N/A'} 
- Score: {top_performer['Score'] if top_performer is not None else 'N/A'}
- Performance Tier: {top_performer['Performance Tier'] if top_performer is not None else 'N/A'}

**Lowest Performer:** {bottom_performer['Supplier'] if bottom_performer is not None else 'N/A'}
- Score: {bottom_performer['Score'] if bottom_performer is not None else 'N/A'}
- Performance Tier: {bottom_performer['Performance Tier'] if bottom_performer is not None else 'N/A'}

**Audit Findings Summary:**
{len([k for k, v in state['audit_findings'].items() if v > 0])} suppliers have audit findings that require attention.

**Recommendation:** Focus on improving suppliers with scores below 70 and address audit findings promptly.
            """.strip()
            
            state["summary"] = fallback_summary
        
        return state

    def plot_graph(self, state: SupplierState) -> SupplierState:
        """Create performance chart"""
        try:
            scores = state["score_table"]
            
            # Use utility to create chart
            chart_path = ChartGenerator.create_performance_chart(
                scores, 
                title="Supplier Performance Scores",
                figsize=(10, 6)
            )
            
            state["graph_path"] = chart_path
            
        except Exception as e:
            print(f"Error creating chart: {str(e)}")
            state["graph_path"] = ""
        
        return state

    def final_output(self, state: SupplierState) -> SupplierState:
        """Final processing - just return the state"""
        return state

    def save_to_doc(self, state: SupplierState) -> bytes:
        """Generate Word document and return as bytes"""
        try:
            doc = Document()
            doc.add_heading("Supplier Performance Report", 0)
            
            # Executive Summary
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(
                "This report evaluates supplier performance using standard industry practices. "
                "The scores reflect metrics including On-Time Delivery, Defect Rates, and Audit Observations. "
                "Weightage recommendations and summary were generated using an AI-powered analysis system."
            )

            # Performance Summary
            doc.add_heading("AI-Generated Performance Summary", level=1)
            doc.add_paragraph(state["summary"])

            # Detailed Scores
            doc.add_heading("Supplier Performance Scores", level=1)
            
            # Use utility to format scores
            scores_df = ReportFormatter.format_score_table(state["score_table"])
            
            # Add table to document
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Rank'
            hdr_cells[1].text = 'Supplier'
            hdr_cells[2].text = 'Score'
            hdr_cells[3].text = 'Performance Tier'
            
            for _, row in scores_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['Rank'])
                row_cells[1].text = str(row['Supplier'])
                row_cells[2].text = str(row['Score'])
                row_cells[3].text = str(row['Performance Tier'])

            # Audit Findings
            if state["audit_findings"]:
                doc.add_heading("Audit Findings Summary", level=1)
                findings_df = ReportFormatter.format_findings_summary(state["audit_findings"])
                
                findings_table = doc.add_table(rows=1, cols=3)
                findings_table.style = 'Table Grid'
                hdr_cells = findings_table.rows[0].cells
                hdr_cells[0].text = 'Supplier'
                hdr_cells[1].text = 'Findings Count'
                hdr_cells[2].text = 'Risk Level'
                
                for _, row in findings_df.iterrows():
                    row_cells = findings_table.add_row().cells
                    row_cells[0].text = str(row['Supplier'])
                    row_cells[1].text = str(row['Findings Count'])
                    row_cells[2].text = str(row['Risk Level'])

            # Weights Used
            if state["weights"]:
                doc.add_heading("Performance Weights Used", level=1)
                weights_para = doc.add_paragraph()
                for metric, weight in state["weights"].items():
                    weights_para.add_run(f"{metric}: {weight:.1%}\n")

            # Add chart if available
            if state["graph_path"] and os.path.exists(state["graph_path"]):
                doc.add_heading("Performance Chart", level=1)
                doc.add_picture(state["graph_path"], width=Inches(5.5))

            # Footer
            doc.add_heading("Report Generation Details", level=1)
            doc.add_paragraph(f"This report was generated using an AI-powered supplier management system.")
            doc.add_paragraph(f"Analysis includes data validation, automated scoring, and AI-generated insights.")

            # Save to bytes
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.getvalue()
            
        except Exception as e:
            print(f"Error creating document: {str(e)}")
            return b""

    def get_detailed_metrics(self, state: SupplierState) -> dict:
        """
        Get detailed metrics for dashboard display
        """
        try:
            scores = state["score_table"]
            findings = state["audit_findings"]
            
            if not scores:
                return {}
            
            # Calculate summary statistics
            score_values = list(scores.values())
            avg_score = sum(score_values) / len(score_values)
            
            # Performance distribution
            tiers = {}
            for score in score_values:
                tier = ScoreCalculator.calculate_performance_tier(score)
                tiers[tier] = tiers.get(tier, 0) + 1
            
            # Risk analysis
            high_risk_suppliers = [name for name, count in findings.items() if count > 2]
            
            return {
                "total_suppliers": len(scores),
                "average_score": round(avg_score, 2),
                "top_performer": max(scores, key=scores.get),
                "top_score": max(score_values),
                "bottom_performer": min(scores, key=scores.get),
                "bottom_score": min(score_values),
                "performance_distribution": tiers,
                "high_risk_suppliers": high_risk_suppliers,
                "total_findings": sum(findings.values()),
                "suppliers_with_findings": len([c for c in findings.values() if c > 0])
            }
            
        except Exception as e:
            print(f"Error calculating detailed metrics: {str(e)}")
            return {}
